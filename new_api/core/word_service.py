"""
Word Document Service - Generate formatted Word documents from structured elements.
"""
import io
import os
import re
from typing import Dict, Any, List, Optional, Union
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_font(
    run,
    font_name: str,
    size: Optional[float] = None,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None
) -> None:
    """
    Set font properties for a run.

    Args:
        run: docx text run
        font_name: Font family name
        size: Font size in points
        bold: Whether to bold the text
        italic: Whether to italicize the text
        color: Font color
    """
    if run is None:
        return

    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    if size is not None:
        run.font.size = Pt(size)

    run.font.bold = bold if bold else None
    run.font.italic = italic if italic else None

    if color is not None:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)


def apply_paragraph_style(p, elem: Dict[str, Any]) -> None:
    """
    Apply paragraph styling from element dictionary.

    Args:
        p: docx paragraph
        elem: Element dictionary with style properties
    """
    alignment = elem.get("alignment")
    if alignment == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alignment == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    spacing_before = elem.get("spacing_before")
    if spacing_before:
        p.paragraph_format.space_before = Pt(spacing_before)

    spacing_after = elem.get("spacing_after")
    if spacing_after:
        p.paragraph_format.space_after = Pt(spacing_after)

    line_spacing = elem.get("line_spacing")
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing

    # First line indent
    if elem.get("indent"):
        p.paragraph_format.first_line_indent = Cm(0.74)

    left_indent = elem.get("left_indent")
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)

    right_indent = elem.get("right_indent")
    if right_indent:
        p.paragraph_format.right_indent = Cm(right_indent)

    # Hanging indent
    if elem.get("hanging_indent"):
        p.paragraph_format.first_line_indent = -Cm(0.74)
        p.paragraph_format.left_indent = Cm(0.74)


def parse_rules_for_defaults(rules: Optional[str] = None) -> Dict[str, Any]:
    """
    Parse formatting rules to extract default font and size.

    Args:
        rules: Formatting rules string

    Returns:
        Dictionary with default body and title styles
    """
    defaults = {
        "body_font": "宋体",
        "body_size": 12,
        "title_font": "黑体",
        "title_size": 22,
    }

    if not rules:
        return defaults

    # Extract body font
    font_match = re.search(r'正文[：:]\s*(\S+)', rules)
    if font_match:
        defaults["body_font"] = font_match.group(1)

    # Extract body size
    size_match = re.search(r'正文[^\n]*?(小四|五号|小五|四号|14|12|10\.5)号?', rules)
    if size_match:
        size_text = size_match.group(1)
        size_map = {"小四": 12, "五号": 10.5, "小五": 9, "四号": 14}
        if size_text in size_map:
            defaults["body_size"] = size_map[size_text]
        else:
            try:
                defaults["body_size"] = float(size_text)
            except ValueError:
                pass

    # Extract title font
    title_match = re.search(r'标题[：:]\s*(\S+)', rules)
    if title_match:
        defaults["title_font"] = title_match.group(1)

    return defaults


class WordDocumentGenerator:
    """
    Generator for Word documents from structured document elements.
    """

    def __init__(self, rules: Optional[str] = None):
        """
        Initialize the document generator.

        Args:
            rules: Formatting rules string
        """
        self.rules = rules
        self.defaults = parse_rules_for_defaults(rules)

    def create_document(self, elements: List[Dict[str, Any]]) -> Document:
        """
        Create a Word document from element list.

        Args:
            elements: List of structured document elements

        Returns:
            docx Document object
        """
        doc = Document()

        # Set default styles
        default_style = doc.styles['Normal']
        default_style.font.name = self.defaults["body_font"]
        if default_style.element.rPr is not None:
            default_style.element.rPr.rFonts.set(qn('w:eastAsia'), self.defaults["body_font"])
        default_style.font.size = Pt(self.defaults["body_size"])

        for elem in elements:
            try:
                self._render_element(doc, elem)
            except Exception as e:
                print(f"[WARNING] Failed to render element: {elem.get('type')}, error: {e}")
                continue

        return doc

    def _render_element(self, doc: Document, elem: Dict[str, Any]) -> None:
        """
        Render a single document element.

        Args:
            doc: docx Document
            elem: Element dictionary
        """
        elem_type = elem.get("type")

        if elem_type == "heading":
            self._render_heading(doc, elem)
        elif elem_type == "paragraph":
            self._render_paragraph(doc, elem)
        elif elem_type == "list":
            self._render_list(doc, elem)
        elif elem_type == "table":
            self._render_table(doc, elem)
        elif elem_type == "formula":
            self._render_formula(doc, elem)
        elif elem_type == "quote":
            self._render_quote(doc, elem)
        elif elem_type == "abstract":
            self._render_abstract(doc, elem)
        elif elem_type == "abstract_en":
            self._render_abstract_en(doc, elem)
        elif elem_type == "keywords":
            self._render_keywords(doc, elem)
        elif elem_type == "keywords_en":
            self._render_keywords_en(doc, elem)
        elif elem_type == "references":
            self._render_references(doc, elem)
        elif elem_type == "author_info":
            self._render_author_info(doc, elem)
        elif elem_type == "code_block":
            self._render_code_block(doc, elem)

    def _render_heading(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render heading element"""
        level = elem.get("level", 1)
        text = elem.get("text", "")

        # Skip rendering these headings - they are handled by their respective elements
        skip_headings = ["关键词", "Keywords", "Key words", "摘要", "Abstract", "参考文献"]
        if text in skip_headings:
            return

        font = elem.get("font", self.defaults["title_font"])
        size = elem.get("size", 22 if level == 1 else (16 if level == 2 else 14))
        is_bold = elem.get("bold", True)

        # For level 4, use regular paragraph to avoid built-in style's italic
        if level == 4:
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font(run, font, size, is_bold, italic=False, color=RGBColor(0, 0, 0))
            p.paragraph_format.left_indent = Cm(0.74)
        else:
            p = doc.add_heading(text, level)
            for run in p.runs:
                set_font(run, font, size, is_bold, italic=False, color=RGBColor(0, 0, 0))

        # Alignment
        alignment = elem.get("alignment")
        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_paragraph(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render paragraph element"""
        text = elem.get("text", "")
        p = doc.add_paragraph(text)

        apply_paragraph_style(p, elem)

        font = elem.get("font", self.defaults["body_font"])
        size = elem.get("size", self.defaults["body_size"])

        if p.runs:
            for run in p.runs:
                set_font(run, font, size, elem.get("bold", False), italic=False, color=RGBColor(0, 0, 0))

    def _render_list(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render list element"""
        items = elem.get("items", [])
        list_style = elem.get("style", "unordered")

        font = elem.get("font", self.defaults["body_font"])
        size = elem.get("size", self.defaults["body_size"])

        for i, item in enumerate(items):
            # Manually add numbering to avoid inheriting from previous lists
            if list_style == "ordered":
                number = f"{i + 1}. "
                p = doc.add_paragraph()
                run = p.add_run(number + item)
            else:
                bullet = elem.get("bullet", "•")
                p = doc.add_paragraph()
                run = p.add_run(f"{bullet} {item}")

            set_font(run, font, size, False, italic=False, color=RGBColor(0, 0, 0))

            if elem.get("hanging_indent"):
                p.paragraph_format.first_line_indent = -Cm(0.74)
                p.paragraph_format.left_indent = Cm(0.74)

    def _render_table(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render table element"""
        headers = elem.get("headers", [])
        rows = elem.get("rows", [])

        if not rows and not headers:
            return

        all_rows = []
        if headers:
            all_rows.append(headers)
        all_rows.extend(rows)

        if not all_rows:
            return

        cols = len(all_rows[0]) if all_rows else 0
        table = doc.add_table(rows=len(all_rows), cols=cols)
        table.style = 'Table Grid'

        is_three_line = elem.get("style") == "三线表"

        for i, row_data in enumerate(all_rows):
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    cell = table.cell(i, j)
                    cell.text = str(cell_text)

                    cell_para = cell.paragraphs[0]
                    if i == 0 and is_three_line:
                        for run in cell_para.runs:
                            set_font(run, '黑体', 10.5, True, italic=False, color=RGBColor(0, 0, 0))
                    else:
                        for run in cell_para.runs:
                            set_font(run, '宋体', 10.5, False, italic=False, color=RGBColor(0, 0, 0))

    def _render_formula(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render formula element"""
        text = elem.get("text", "")
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_font(run, '等宽字体', None, False, italic=False, color=RGBColor(0, 0, 0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _render_quote(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render quote element"""
        text = elem.get("text", "")
        speaker = elem.get("speaker", "")
        p = doc.add_paragraph()
        run = p.add_run(text)

        set_font(run, elem.get("font", "楷体"), elem.get("size", 12), False, italic=False, color=RGBColor(0, 0, 0))

        if speaker:
            p.add_run(f"\n—— {speaker}")

        left = elem.get("left_indent")
        right = elem.get("right_indent")
        if left is not None:
            p.paragraph_format.left_indent = Cm(left)
        if right is not None:
            p.paragraph_format.right_indent = Cm(right)

    def _render_abstract(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render Chinese abstract element"""
        doc.add_heading("摘要", level=1)
        p = doc.add_paragraph(elem.get("text", ""))
        apply_paragraph_style(p, elem)

        if p.runs:
            for run in p.runs:
                set_font(
                    run,
                    elem.get("font", self.defaults["body_font"]),
                    elem.get("size", self.defaults["body_size"]),
                    elem.get("bold", False),
                    italic=False,
                    color=RGBColor(0, 0, 0)
                )

    def _render_abstract_en(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render English abstract element"""
        doc.add_heading("Abstract", level=1)
        p = doc.add_paragraph(elem.get("text", ""))
        apply_paragraph_style(p, elem)

        if p.runs:
            for run in p.runs:
                set_font(
                    run,
                    elem.get("font", "Times New Roman"),
                    elem.get("size", self.defaults["body_size"]),
                    elem.get("bold", False),
                    italic=False,
                    color=RGBColor(0, 0, 0)
                )

    def _render_keywords(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render Chinese keywords element"""
        keywords = elem.get("keywords", [])
        if isinstance(keywords, str):
            keywords = [k.strip() for k in keywords.replace('；', ',').replace(';', ',').split(',')]
            keywords = [k for k in keywords if k]

        p = doc.add_paragraph("关键词：" + "；".join(keywords))
        apply_paragraph_style(p, elem)

        if p.runs:
            for run in p.runs:
                set_font(
                    run,
                    elem.get("font", self.defaults["body_font"]),
                    elem.get("size", self.defaults["body_size"]),
                    elem.get("bold", False),
                    italic=False,
                    color=RGBColor(0, 0, 0)
                )

    def _render_keywords_en(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render English keywords element"""
        keywords = elem.get("keywords", [])
        if isinstance(keywords, str):
            keywords = [k.strip() for k in keywords.split(';')]

        p = doc.add_paragraph("Key words: " + "; ".join(keywords))
        apply_paragraph_style(p, elem)

        if p.runs:
            for run in p.runs:
                set_font(
                    run,
                    elem.get("font", "Times New Roman"),
                    elem.get("size", self.defaults["body_size"]),
                    elem.get("bold", False),
                    italic=False,
                    color=RGBColor(0, 0, 0)
                )

    def _render_references(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render references element"""
        doc.add_heading("参考文献", level=1)
        for ref in elem.get("items", []):
            p = doc.add_paragraph(ref)
            p.paragraph_format.first_line_indent = -Cm(2)
            p.paragraph_format.left_indent = Cm(2)

            if p.runs:
                for run in p.runs:
                    set_font(
                        run,
                        self.defaults["body_font"],
                        self.defaults["body_size"],
                        False,
                        italic=False,
                        color=RGBColor(0, 0, 0)
                    )

    def _render_author_info(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render author info element"""
        authors = elem.get("authors", [])
        if isinstance(authors, str):
            authors = [a.strip() for a in authors.split(',')]

        p = doc.add_paragraph(", ".join(authors))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if p.runs:
            for run in p.runs:
                set_font(
                    run,
                    self.defaults["body_font"],
                    self.defaults["body_size"],
                    False,
                    italic=False,
                    color=RGBColor(0, 0, 0)
                )

    def _render_code_block(self, doc: Document, elem: Dict[str, Any]) -> None:
        """Render code block element with monospace font"""
        text = elem.get("text", "")

        # Use monospace font (等宽字体 / Courier New)
        font = elem.get("font", "等宽字体")
        size = elem.get("size", 10.5)

        # Add left indentation for code blocks
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Cm(1)

        if p.runs:
            for run in p.runs:
                set_font(run, font, size, False, italic=False, color=RGBColor(0, 0, 0))


# Convenience functions
def create_word_document(
    elements: List[Dict[str, Any]],
    rules: Optional[str] = None
) -> Document:
    """
    Create a Word document from element list.

    Args:
        elements: List of structured document elements
        rules: Formatting rules string

    Returns:
        docx Document object
    """
    generator = WordDocumentGenerator(rules)
    return generator.create_document(elements)


def save_document(
    doc: Document,
    output_path: Optional[str] = None
) -> Union[bool, io.BytesIO]:
    """
    Save Word document to file or return as BytesIO.

    Args:
        doc: docx Document object
        output_path: Path to save the document (optional)

    Returns:
        If output_path provided: True if successful, False otherwise
        If no output_path: Returns BytesIO object with the document
    """
    try:
        if output_path:
            # Save to file (local development)
            abs_dir = os.path.dirname(os.path.abspath(output_path))
            os.makedirs(abs_dir, exist_ok=True)
            doc.save(output_path)
            return os.path.exists(output_path)
        else:
            # Return as BytesIO (cloud deployment)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
    except Exception as e:
        print(f"[ERROR] Failed to save document: {e}")
        return False
