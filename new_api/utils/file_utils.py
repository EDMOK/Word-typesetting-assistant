"""
Utility functions for file operations.
"""
import os
import tempfile
import logging
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document, including paragraphs and tables"""
    try:
        doc = Document(file_path)
        text_parts = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))

        if not text_parts:
            raise ValueError("文档中未提取到任何文本内容")

        return '\n'.join(text_parts)
    except ValueError:
        # Re-raise ValueError with original message
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Extract text from Word document bytes (for cloud deployment), including paragraphs and tables"""
    tmp_path = None
    try:
        # Create a temporary file since python-docx needs a file path
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name

        doc = Document(tmp_path)
        text_parts = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))

        if not text_parts:
            raise ValueError("文档中未提取到任何文本内容")

        return '\n'.join(text_parts)
    except ValueError:
        # Re-raise ValueError with original message
        raise
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")
    finally:
        # Clean up temp file
        if tmp_path:
            try:
                os.remove(tmp_path)
            except OSError:
                pass


def decode_file_content(content: bytes, encodings=None) -> str:
    """Decode file content with multiple encoding attempts"""
    if encodings is None:
        encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin-1']

    for encoding in encodings:
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue

    return content.decode('utf-8', errors='ignore')


def cleanup_temp_file(file_path: str) -> None:
    """Remove temporary file if exists"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except OSError:
        pass
